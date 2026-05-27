from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


@dataclass
class ParsedDocument:
    text: str
    metadata: dict[str, str]


def parse_file(path: Path, *, excel_max_rows: int = 80) -> ParsedDocument:
    suffix = path.suffix.lower()
    meta = {"source_path": str(path), "file_name": path.name, "extension": suffix}

    if suffix == ".pdf":
        return _parse_pdf(path, meta)
    if suffix == ".docx":
        return _parse_docx(path, meta)
    if suffix in {".txt", ".md", ".html", ".htm", ".rtf", ".csv"}:
        return _parse_text(path, meta)
    if suffix in {".xlsx", ".xls"}:
        return _parse_excel(path, meta, max_rows=excel_max_rows)
    raise ValueError(f"Unsupported extension: {suffix}")


def _parse_pdf(path: Path, meta: dict[str, str]) -> ParsedDocument:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {i}]\n{text}")
    return ParsedDocument("\n\n".join(pages), meta)


def _parse_docx(path: Path, meta: dict[str, str]) -> ParsedDocument:
    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return ParsedDocument("\n\n".join(paras), meta)


def _parse_text(path: Path, meta: dict[str, str]) -> ParsedDocument:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text = path.read_text(encoding=enc)
            return ParsedDocument(text, meta)
        except UnicodeDecodeError:
            continue
    text = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(text, meta)


def _parse_excel(path: Path, meta: dict[str, str], *, max_rows: int) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".xls":
        df_map = pd.read_excel(path, sheet_name=None, engine="xlrd")
    else:
        df_map = pd.read_excel(path, sheet_name=None, engine="openpyxl")

    parts: list[str] = []
    for sheet_name, df in df_map.items():
        if df.empty:
            continue
        total = len(df)
        chunk_count = max(1, (total + max_rows - 1) // max_rows)
        for i in range(chunk_count):
            start = i * max_rows
            end = min(start + max_rows, total)
            sub = df.iloc[start:end]
            header = f"[Sheet: {sheet_name} | rows {start + 1}-{end} of {total}]"
            parts.append(f"{header}\n{sub.to_markdown(index=False)}")
    return ParsedDocument("\n\n".join(parts), {**meta, "format": "excel"})


def parse_file_safe(path: Path, *, excel_max_rows: int = 80) -> ParsedDocument | None:
    try:
        return parse_file(path, excel_max_rows=excel_max_rows)
    except Exception:
        return None
